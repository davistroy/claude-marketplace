#!/usr/bin/env python3
"""Generate a professional feedback assessment .docx document from structured JSON input.

Standalone script â€” no imports from the voice-capture src/ package.

Usage:
    python -m feedback_docx_generator --input data.json --output report.docx
    cat data.json | python -m feedback_docx_generator --output report.docx
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(
        "ERROR: python-docx is not installed.\n"
        "Install it with: pip install python-docx>=1.0",
        file=sys.stderr,
    )
    sys.exit(1)


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------
FONT_BODY = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_HEADING1 = Pt(18)
FONT_SIZE_HEADING2 = Pt(14)
FONT_SIZE_HEADING3 = Pt(12)
FONT_SIZE_METADATA = Pt(9)
COLOR_HEADING = RGBColor(0x1B, 0x3A, 0x5C)  # dark blue
COLOR_METADATA = RGBColor(0x6B, 0x6B, 0x6B)  # grey
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_run_style(run, *, font_name=FONT_BODY, size=FONT_SIZE_BODY,
                   color=COLOR_BODY, bold=False, italic=False):
    """Apply consistent font styling to a run."""
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def _add_heading(doc, text, level=1):
    """Add a styled heading paragraph."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = COLOR_HEADING
        run.font.name = FONT_BODY
        if level == 1:
            run.font.size = FONT_SIZE_HEADING1
        elif level == 2:
            run.font.size = FONT_SIZE_HEADING2
        else:
            run.font.size = FONT_SIZE_HEADING3
    return heading


def _add_body(doc, text):
    """Add a body paragraph with standard styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_style(run)
    return para


def _add_metadata_line(doc, text):
    """Add a grey metadata line."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_style(run, size=FONT_SIZE_METADATA, color=COLOR_METADATA, italic=True)
    return para


def _add_bullet(doc, text, *, bold_prefix=None):
    """Add a bullet-point paragraph, optionally with a bold prefix."""
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = para.add_run(bold_prefix)
        _set_run_style(run_bold, bold=True)
        run_rest = para.add_run(text)
        _set_run_style(run_rest)
    else:
        run = para.add_run(text)
        _set_run_style(run)
    return para


def _format_date(date_str):
    """Best-effort date formatting. Returns the original string on failure."""
    if not date_str:
        return "N/A"
    for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f",
                "%Y-%m-%dT%H:%M:%S%z"):
        try:
            return datetime.strptime(date_str[:19], fmt[:min(len(fmt), 19)]).strftime("%B %d, %Y")
        except (ValueError, TypeError):
            continue
    return date_str


def _safe_get(data, *keys, default=""):
    """Safely traverse nested dicts."""
    current = data
    for key in keys:
        if isinstance(current, dict):
            current = current.get(key, default)
        else:
            return default
    return current if current is not None else default


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _build_title_page(doc, data):
    """Title page header with employee name, period, and metadata."""
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Employee Feedback Assessment")
    _set_run_style(run, size=Pt(24), color=COLOR_HEADING, bold=True)

    # Employee name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(data.get("employee_name", "Unknown Employee"))
    _set_run_style(run, size=Pt(18), color=COLOR_BODY)

    # Spacer
    doc.add_paragraph()

    # Metadata block
    period = data.get("assessment_period", {})
    start = _format_date(period.get("start", ""))
    end = _format_date(period.get("end", ""))
    gen_date = _format_date(data.get("generation_date", ""))
    total = data.get("total_entries", 0)

    meta_lines = [
        f"Assessment Period: {start} to {end}",
        f"Generated: {gen_date}",
        f"Based on {total} feedback {'entry' if total == 1 else 'entries'}",
    ]
    for line in meta_lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        _set_run_style(run, size=FONT_SIZE_METADATA, color=COLOR_METADATA, italic=True)

    # Page break after title section
    doc.add_page_break()


def _build_executive_summary(doc, synthesis):
    """Executive Summary section."""
    _add_heading(doc, "Executive Summary", level=1)
    summary_text = _safe_get(synthesis, "executive_summary", default="No executive summary provided.")
    _add_body(doc, summary_text)


def _build_strengths(doc, synthesis):
    """Strengths section with evidence citations."""
    _add_heading(doc, "Strengths", level=1)
    strengths = synthesis.get("strengths", [])
    if not strengths:
        _add_body(doc, "No strengths identified in the available feedback entries.")
        return

    for i, strength in enumerate(strengths):
        name = strength.get("name", f"Strength {i + 1}")
        description = strength.get("description", "")
        frequency = strength.get("frequency", "")
        evidence = strength.get("evidence", [])

        _add_heading(doc, name, level=2)
        if frequency:
            _add_metadata_line(doc, f"Frequency: {frequency}")
        if description:
            _add_body(doc, description)
        if evidence:
            para = doc.add_paragraph()
            run = para.add_run("Supporting Evidence:")
            _set_run_style(run, bold=True, size=FONT_SIZE_BODY)
            for item in evidence:
                date_str = _format_date(item.get("date", ""))
                summary = item.get("summary", "")
                _add_bullet(doc, f" {summary}", bold_prefix=f"[{date_str}]")


def _build_areas_for_development(doc, synthesis):
    """Areas for Development section with evidence citations."""
    _add_heading(doc, "Areas for Development", level=1)
    areas = synthesis.get("areas_for_development", [])
    if not areas:
        _add_body(doc, "No areas for development identified in the available feedback entries.")
        return

    for i, area in enumerate(areas):
        name = area.get("name", f"Area {i + 1}")
        description = area.get("description", "")
        pattern = area.get("pattern", "")
        evidence = area.get("evidence", [])

        _add_heading(doc, name, level=2)
        if pattern:
            _add_metadata_line(doc, f"Pattern: {pattern}")
        if description:
            _add_body(doc, description)
        if evidence:
            para = doc.add_paragraph()
            run = para.add_run("Supporting Evidence:")
            _set_run_style(run, bold=True, size=FONT_SIZE_BODY)
            for item in evidence:
                date_str = _format_date(item.get("date", ""))
                summary = item.get("summary", "")
                _add_bullet(doc, f" {summary}", bold_prefix=f"[{date_str}]")


def _build_patterns_and_themes(doc, synthesis):
    """Patterns and Themes section."""
    _add_heading(doc, "Patterns and Themes", level=1)
    patterns = synthesis.get("patterns_and_themes", {})
    if not patterns:
        _add_body(doc, "Insufficient data to identify patterns and themes.")
        return

    sections = [
        ("Trends Over Time", "trends"),
        ("Relationships Between Observations", "relationships"),
        ("Situational Patterns", "situational"),
    ]
    for title, key in sections:
        text = patterns.get(key, "")
        if text:
            _add_heading(doc, title, level=2)
            _add_body(doc, text)


def _build_recommendations(doc, synthesis):
    """Recommendations section grouped by type."""
    _add_heading(doc, "Recommendations", level=1)
    recommendations = synthesis.get("recommendations", [])
    if not recommendations:
        _add_body(doc, "No specific recommendations at this time.")
        return

    # Group by type
    grouped = {"Continue": [], "Develop": [], "Stretch": []}
    for rec in recommendations:
        rec_type = rec.get("type", "Develop")
        grouped.setdefault(rec_type, []).append(rec)

    type_descriptions = {
        "Continue": "Behaviors and skills to maintain and reinforce",
        "Develop": "Areas needing improvement with specific development actions",
        "Stretch": "Growth opportunities that build on existing strengths",
    }

    for rec_type in ("Continue", "Develop", "Stretch"):
        items = grouped.get(rec_type, [])
        if not items:
            continue
        _add_heading(doc, rec_type, level=2)
        _add_metadata_line(doc, type_descriptions.get(rec_type, ""))
        for item in items:
            recommendation = item.get("recommendation", "")
            rationale = item.get("rationale", "")
            _add_bullet(doc, recommendation)
            if rationale:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(f"Rationale: {rationale}")
                _set_run_style(run, size=FONT_SIZE_METADATA, color=COLOR_METADATA, italic=True)


def _build_appendix(doc, entries):
    """Appendix with individual feedback entries in chronological order."""
    doc.add_page_break()
    _add_heading(doc, "Appendix: Individual Feedback Entries", level=1)

    if not entries:
        _add_body(doc, "No individual entries to display.")
        return

    # Sort chronologically
    sorted_entries = sorted(entries, key=lambda e: e.get("date", ""))

    for i, entry in enumerate(sorted_entries):
        title = entry.get("title", f"Entry {i + 1}")
        date_str = _format_date(entry.get("date", ""))
        feedback_type = entry.get("feedback_type", "Observation")
        summary = entry.get("summary", "")
        context = entry.get("context", "")
        actionable = entry.get("actionable_items", "")
        transcript = entry.get("raw_transcript", "")
        tags = entry.get("tags", [])

        # Entry heading
        _add_heading(doc, f"{title}", level=2)
        _add_metadata_line(doc, f"{date_str}  |  {feedback_type}")

        if tags:
            tag_str = ", ".join(tags) if isinstance(tags, list) else str(tags)
            _add_metadata_line(doc, f"Tags: {tag_str}")

        if summary:
            para = doc.add_paragraph()
            run_label = para.add_run("Summary: ")
            _set_run_style(run_label, bold=True)
            run_text = para.add_run(summary)
            _set_run_style(run_text)

        if context:
            para = doc.add_paragraph()
            run_label = para.add_run("Context: ")
            _set_run_style(run_label, bold=True)
            run_text = para.add_run(context)
            _set_run_style(run_text)

        if actionable:
            para = doc.add_paragraph()
            run_label = para.add_run("Actionable Items: ")
            _set_run_style(run_label, bold=True)
            run_text = para.add_run(actionable)
            _set_run_style(run_text)

        if transcript:
            _add_heading(doc, "Raw Transcript", level=3)
            para = doc.add_paragraph()
            run = para.add_run(transcript)
            _set_run_style(run, size=Pt(10), color=COLOR_METADATA)

        # Separator between entries (except last)
        if i < len(sorted_entries) - 1:
            separator = doc.add_paragraph()
            run = separator.add_run("_" * 60)
            _set_run_style(run, color=RGBColor(0xCC, 0xCC, 0xCC), size=Pt(8))


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------

def generate_docx(data: dict, output_path: str):
    """Build the complete .docx document from structured JSON data."""
    doc = Document()

    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = FONT_SIZE_BODY
    font.color.rgb = COLOR_BODY

    synthesis = data.get("synthesis", {})
    entries = data.get("entries", [])

    _build_title_page(doc, data)
    _build_executive_summary(doc, synthesis)
    _build_strengths(doc, synthesis)
    _build_areas_for_development(doc, synthesis)
    _build_patterns_and_themes(doc, synthesis)
    _build_recommendations(doc, synthesis)
    _build_appendix(doc, entries)

    # Footer metadata
    doc.add_paragraph()
    _add_metadata_line(
        doc,
        f"Generated by Feedback Synthesis | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    )

    # Ensure output directory exists
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output))
    return str(output.resolve())


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate a feedback assessment .docx from JSON input.",
    )
    parser.add_argument(
        "--input",
        type=str,
        default=None,
        help="Path to input JSON file. Reads from stdin if omitted.",
    )
    parser.add_argument(
        "--output",
        type=str,
        required=True,
        help="Path for the output .docx file.",
    )
    args = parser.parse_args()

    # Read JSON input
    if args.input:
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"ERROR: Input file not found: {args.input}", file=sys.stderr)
            sys.exit(1)
        with open(input_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.load(sys.stdin)

    # Generate document
    result_path = generate_docx(data, args.output)
    print(f"Document generated: {result_path}")


if __name__ == "__main__":
    main()
