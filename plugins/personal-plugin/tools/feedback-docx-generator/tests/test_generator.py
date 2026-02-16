"""Tests for the document generation functions in __main__.py."""

from docx import Document
from docx.shared import Pt, RGBColor

from feedback_docx_generator.__main__ import (
    COLOR_BODY,
    COLOR_METADATA,
    FONT_BODY,
    FONT_SIZE_BODY,
    _add_body,
    _add_bullet,
    _add_heading,
    _add_metadata_line,
    _build_appendix,
    _build_areas_for_development,
    _build_executive_summary,
    _build_patterns_and_themes,
    _build_recommendations,
    _build_strengths,
    _build_title_page,
    _format_date,
    _safe_get,
    _set_run_style,
    generate_docx,
)


# ---------------------------------------------------------------------------
# Helper function tests
# ---------------------------------------------------------------------------


class TestFormatDate:
    """Tests for _format_date."""

    def test_iso_date(self):
        """Test standard ISO date format."""
        result = _format_date("2025-07-15")
        assert result == "July 15, 2025"

    def test_iso_datetime(self):
        """Test ISO datetime format."""
        result = _format_date("2025-07-15T14:30:00")
        assert result == "July 15, 2025"

    def test_iso_datetime_with_microseconds(self):
        """Test ISO datetime with microseconds."""
        result = _format_date("2025-07-15T14:30:00.123456")
        assert result == "July 15, 2025"

    def test_empty_string_returns_na(self):
        """Test that empty string returns N/A."""
        assert _format_date("") == "N/A"

    def test_none_returns_na(self):
        """Test that None returns N/A."""
        assert _format_date(None) == "N/A"

    def test_unparseable_returns_original(self):
        """Test that unparseable strings are returned as-is."""
        assert _format_date("not-a-date") == "not-a-date"

    def test_partial_date_returns_original(self):
        """Test that partial date strings are returned as-is."""
        assert _format_date("2025-13-45") == "2025-13-45"


class TestSafeGet:
    """Tests for _safe_get."""

    def test_single_key(self):
        """Test retrieving a single key."""
        assert _safe_get({"a": "value"}, "a") == "value"

    def test_nested_keys(self):
        """Test retrieving nested keys."""
        data = {"a": {"b": {"c": "deep"}}}
        assert _safe_get(data, "a", "b", "c") == "deep"

    def test_missing_key_returns_default(self):
        """Test missing key returns default."""
        assert _safe_get({"a": 1}, "b") == ""

    def test_custom_default(self):
        """Test custom default value."""
        assert _safe_get({}, "missing", default="N/A") == "N/A"

    def test_none_value_returns_default(self):
        """Test that None value returns default."""
        assert _safe_get({"a": None}, "a") == ""

    def test_non_dict_intermediate_returns_default(self):
        """Test non-dict intermediate returns default."""
        assert _safe_get({"a": "string"}, "a", "b") == ""


class TestSetRunStyle:
    """Tests for _set_run_style."""

    def test_default_style(self):
        """Test default style application."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Test")
        _set_run_style(run)

        assert run.font.name == FONT_BODY
        assert run.font.size == FONT_SIZE_BODY
        assert run.font.color.rgb == COLOR_BODY
        assert run.font.bold is not True
        assert run.font.italic is not True

    def test_bold_style(self):
        """Test bold style application."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold")
        _set_run_style(run, bold=True)

        assert run.font.bold is True

    def test_custom_color_and_size(self):
        """Test custom color and size."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Custom")
        custom_color = RGBColor(0xFF, 0x00, 0x00)
        custom_size = Pt(20)
        _set_run_style(run, color=custom_color, size=custom_size)

        assert run.font.color.rgb == custom_color
        assert run.font.size == custom_size


# ---------------------------------------------------------------------------
# Document element tests
# ---------------------------------------------------------------------------


class TestAddHeading:
    """Tests for _add_heading."""

    def test_level_1_heading(self):
        """Test level 1 heading creation."""
        doc = Document()
        heading = _add_heading(doc, "Test Heading", level=1)
        assert heading is not None
        # Verify the text content
        text = heading.text
        assert "Test Heading" in text

    def test_level_2_heading(self):
        """Test level 2 heading creation."""
        doc = Document()
        heading = _add_heading(doc, "Sub Heading", level=2)
        assert heading is not None

    def test_level_3_heading(self):
        """Test level 3 heading creation."""
        doc = Document()
        heading = _add_heading(doc, "Sub-sub Heading", level=3)
        assert heading is not None


class TestAddBody:
    """Tests for _add_body."""

    def test_adds_paragraph(self):
        """Test that body paragraph is added."""
        doc = Document()
        para = _add_body(doc, "Body text here")
        assert para is not None
        assert para.runs[0].text == "Body text here"


class TestAddMetadataLine:
    """Tests for _add_metadata_line."""

    def test_metadata_style(self):
        """Test metadata line styling."""
        doc = Document()
        para = _add_metadata_line(doc, "Generated: 2025-01-01")
        run = para.runs[0]
        assert run.font.italic is True
        assert run.font.color.rgb == COLOR_METADATA


class TestAddBullet:
    """Tests for _add_bullet."""

    def test_simple_bullet(self):
        """Test simple bullet without bold prefix."""
        doc = Document()
        para = _add_bullet(doc, "Bullet text")
        assert para.runs[0].text == "Bullet text"

    def test_bullet_with_bold_prefix(self):
        """Test bullet with bold prefix."""
        doc = Document()
        para = _add_bullet(doc, " rest of text", bold_prefix="[2025-01-01]")
        assert para.runs[0].text == "[2025-01-01]"
        assert para.runs[0].font.bold is True
        assert para.runs[1].text == " rest of text"


# ---------------------------------------------------------------------------
# Section builder tests
# ---------------------------------------------------------------------------


class TestBuildTitlePage:
    """Tests for _build_title_page."""

    def test_title_page_content(self, sample_feedback_data):
        """Test that title page contains expected elements."""
        doc = Document()
        _build_title_page(doc, sample_feedback_data)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Employee Feedback Assessment" in all_text
        assert "Jane Doe" in all_text

    def test_title_page_metadata(self, sample_feedback_data):
        """Test that title page shows metadata."""
        doc = Document()
        _build_title_page(doc, sample_feedback_data)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "5 feedback entries" in all_text

    def test_title_page_single_entry(self):
        """Test singular 'entry' for total_entries=1."""
        doc = Document()
        data = {"employee_name": "Test", "total_entries": 1}
        _build_title_page(doc, data)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "1 feedback entry" in all_text

    def test_title_page_missing_name(self):
        """Test default name when employee_name is missing."""
        doc = Document()
        _build_title_page(doc, {})

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Unknown Employee" in all_text


class TestBuildExecutiveSummary:
    """Tests for _build_executive_summary."""

    def test_with_summary(self, sample_feedback_data):
        """Test executive summary with content."""
        doc = Document()
        _build_executive_summary(doc, sample_feedback_data["synthesis"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Executive Summary" in all_text
        assert "strong leadership" in all_text

    def test_without_summary(self):
        """Test executive summary when missing."""
        doc = Document()
        _build_executive_summary(doc, {})

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No executive summary provided." in all_text


class TestBuildStrengths:
    """Tests for _build_strengths."""

    def test_with_strengths(self, sample_feedback_data):
        """Test strengths section with data."""
        doc = Document()
        _build_strengths(doc, sample_feedback_data["synthesis"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Strengths" in all_text
        assert "Technical Leadership" in all_text
        assert "Communication" in all_text

    def test_strength_evidence(self, sample_feedback_data):
        """Test that strength evidence is rendered."""
        doc = Document()
        _build_strengths(doc, sample_feedback_data["synthesis"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Supporting Evidence" in all_text
        assert "microservices" in all_text

    def test_no_strengths(self):
        """Test empty strengths message."""
        doc = Document()
        _build_strengths(doc, {"strengths": []})

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No strengths identified" in all_text

    def test_strength_without_name_uses_default(self):
        """Test that unnamed strengths get default names."""
        doc = Document()
        synthesis = {"strengths": [{"description": "Some strength"}]}
        _build_strengths(doc, synthesis)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Strength 1" in all_text


class TestBuildAreasForDevelopment:
    """Tests for _build_areas_for_development."""

    def test_with_areas(self, sample_feedback_data):
        """Test areas for development with data."""
        doc = Document()
        _build_areas_for_development(doc, sample_feedback_data["synthesis"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Areas for Development" in all_text
        assert "Time Management" in all_text

    def test_no_areas(self):
        """Test empty areas message."""
        doc = Document()
        _build_areas_for_development(doc, {"areas_for_development": []})

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No areas for development identified" in all_text


class TestBuildPatternsAndThemes:
    """Tests for _build_patterns_and_themes."""

    def test_with_patterns(self, sample_feedback_data):
        """Test patterns section with data."""
        doc = Document()
        _build_patterns_and_themes(doc, sample_feedback_data["synthesis"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Patterns and Themes" in all_text
        assert "Trends Over Time" in all_text
        assert "delegation" in all_text

    def test_no_patterns(self):
        """Test empty patterns message."""
        doc = Document()
        _build_patterns_and_themes(doc, {"patterns_and_themes": {}})

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Insufficient data" in all_text


class TestBuildRecommendations:
    """Tests for _build_recommendations."""

    def test_all_recommendation_types(self, sample_feedback_data):
        """Test that all recommendation types are rendered."""
        doc = Document()
        _build_recommendations(doc, sample_feedback_data["synthesis"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Continue" in all_text
        assert "Develop" in all_text
        assert "Stretch" in all_text

    def test_recommendation_rationale(self, sample_feedback_data):
        """Test that rationale is included."""
        doc = Document()
        _build_recommendations(doc, sample_feedback_data["synthesis"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Rationale" in all_text
        assert "track record" in all_text

    def test_no_recommendations(self):
        """Test empty recommendations message."""
        doc = Document()
        _build_recommendations(doc, {"recommendations": []})

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No specific recommendations" in all_text

    def test_unknown_recommendation_type_excluded(self):
        """Test that unknown types are not rendered (only Continue/Develop/Stretch)."""
        doc = Document()
        synthesis = {
            "recommendations": [
                {
                    "type": "CustomType",
                    "recommendation": "Do something custom.",
                }
            ]
        }
        _build_recommendations(doc, synthesis)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Unknown types are not in the iteration order, so they are not rendered
        assert "Do something custom." not in all_text

    def test_default_type_is_develop(self):
        """Test that recommendations without type default to Develop."""
        doc = Document()
        synthesis = {
            "recommendations": [
                {
                    "recommendation": "Improve testing.",
                }
            ]
        }
        _build_recommendations(doc, synthesis)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Develop" in all_text
        assert "Improve testing." in all_text


class TestBuildAppendix:
    """Tests for _build_appendix."""

    def test_appendix_entries(self, sample_feedback_data):
        """Test that appendix entries are rendered."""
        doc = Document()
        _build_appendix(doc, sample_feedback_data["entries"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Appendix" in all_text
        assert "Architecture Review" in all_text
        assert "Sprint Retrospective" in all_text

    def test_appendix_sorted_chronologically(self, sample_feedback_data):
        """Test that entries are sorted by date."""
        doc = Document()
        _build_appendix(doc, sample_feedback_data["entries"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Architecture Review (2025-02-15) should appear before Sprint Retro (2025-05-01)
        arch_pos = all_text.index("Architecture Review")
        sprint_pos = all_text.index("Sprint Retrospective")
        assert arch_pos < sprint_pos

    def test_appendix_no_entries(self):
        """Test empty entries message."""
        doc = Document()
        _build_appendix(doc, [])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No individual entries to display." in all_text

    def test_appendix_entry_details(self, sample_feedback_data):
        """Test that entry details are rendered."""
        doc = Document()
        _build_appendix(doc, sample_feedback_data["entries"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Summary:" in all_text
        assert "Context:" in all_text
        assert "Actionable Items:" in all_text

    def test_appendix_tags(self, sample_feedback_data):
        """Test that tags are rendered."""
        doc = Document()
        _build_appendix(doc, sample_feedback_data["entries"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "leadership" in all_text
        assert "architecture" in all_text

    def test_appendix_raw_transcript(self, sample_feedback_data):
        """Test that raw transcript is rendered."""
        doc = Document()
        _build_appendix(doc, sample_feedback_data["entries"])

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Raw Transcript" in all_text
        assert "Jane presented" in all_text

    def test_appendix_tags_as_string(self):
        """Test that tags provided as string are handled."""
        doc = Document()
        entries = [{"title": "Test", "tags": "single-tag", "date": "2025-01-01"}]
        _build_appendix(doc, entries)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "single-tag" in all_text

    def test_appendix_separator_not_after_last(self, sample_feedback_data):
        """Test that separator line count is len(entries)-1."""
        doc = Document()
        _build_appendix(doc, sample_feedback_data["entries"])

        # Count separator paragraphs (runs with "_" * 60)
        separator_count = sum(
            1 for p in doc.paragraphs if p.runs and "_" * 60 in p.runs[0].text
        )
        assert separator_count == len(sample_feedback_data["entries"]) - 1


# ---------------------------------------------------------------------------
# End-to-end generate_docx tests
# ---------------------------------------------------------------------------


class TestGenerateDocx:
    """Tests for generate_docx."""

    def test_generates_valid_docx(self, sample_feedback_data, temp_output_dir):
        """Test that a valid .docx file is created."""
        output_path = str(temp_output_dir / "test_report.docx")
        result = generate_docx(sample_feedback_data, output_path)

        assert result.endswith("test_report.docx")
        # Verify the file is a valid docx
        doc = Document(result)
        assert len(doc.paragraphs) > 0

    def test_creates_output_directory(self, sample_feedback_data, tmp_path):
        """Test that parent directories are created if needed."""
        output_path = str(tmp_path / "nested" / "dir" / "report.docx")
        result = generate_docx(sample_feedback_data, output_path)
        assert result.endswith("report.docx")

    def test_empty_data_produces_docx(self, empty_feedback_data, temp_output_dir):
        """Test that empty data still produces a valid document."""
        output_path = str(temp_output_dir / "empty_report.docx")
        result = generate_docx(empty_feedback_data, output_path)

        doc = Document(result)
        assert len(doc.paragraphs) > 0

    def test_all_major_sections_present(self, sample_feedback_data, temp_output_dir):
        """Test that all major document sections are present."""
        output_path = str(temp_output_dir / "sections_report.docx")
        result = generate_docx(sample_feedback_data, output_path)
        doc = Document(result)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        expected_sections = [
            "Employee Feedback Assessment",
            "Executive Summary",
            "Strengths",
            "Areas for Development",
            "Patterns and Themes",
            "Recommendations",
            "Appendix",
        ]
        for section in expected_sections:
            assert section in all_text, f"Missing section: {section}"

    def test_footer_metadata(self, sample_feedback_data, temp_output_dir):
        """Test that footer metadata line is present."""
        output_path = str(temp_output_dir / "footer_report.docx")
        result = generate_docx(sample_feedback_data, output_path)
        doc = Document(result)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Generated by Feedback Synthesis" in all_text

    def test_default_document_style(self, sample_feedback_data, temp_output_dir):
        """Test that default Normal style is set correctly."""
        output_path = str(temp_output_dir / "style_report.docx")
        result = generate_docx(sample_feedback_data, output_path)
        doc = Document(result)

        normal_style = doc.styles["Normal"]
        assert normal_style.font.name == FONT_BODY
        assert normal_style.font.size == FONT_SIZE_BODY
        assert normal_style.font.color.rgb == COLOR_BODY
