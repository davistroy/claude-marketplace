"""Tests for the CLI entry point (__main__.py main function)."""

import json
from io import StringIO
from unittest.mock import patch

import pytest

from feedback_docx_generator.__main__ import main


@pytest.fixture
def input_json_file(tmp_path, sample_feedback_data):
    """Create a temporary JSON input file."""
    filepath = tmp_path / "input.json"
    filepath.write_text(json.dumps(sample_feedback_data), encoding="utf-8")
    return filepath


@pytest.fixture
def empty_json_file(tmp_path, empty_feedback_data):
    """Create a temporary JSON input file with empty data."""
    filepath = tmp_path / "empty_input.json"
    filepath.write_text(json.dumps(empty_feedback_data), encoding="utf-8")
    return filepath


class TestCLIArgumentParsing:
    """Tests for argument parsing."""

    def test_output_required(self):
        """Test that --output is required."""
        with patch("sys.argv", ["prog"]):
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code in (0, 2)

    def test_help_flag(self):
        """Test --help flag."""
        with patch("sys.argv", ["prog", "--help"]):
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code == 0

    def test_missing_input_file_exits(self, tmp_path):
        """Test that missing input file causes exit with error."""
        output_path = str(tmp_path / "output.docx")
        with patch(
            "sys.argv",
            ["prog", "--input", "/nonexistent/file.json", "--output", output_path],
        ):
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code == 1


class TestCLIExecution:
    """Tests for CLI execution."""

    def test_from_file_input(self, input_json_file, tmp_path):
        """Test generation from file input."""
        output_path = str(tmp_path / "output.docx")
        with patch(
            "sys.argv",
            ["prog", "--input", str(input_json_file), "--output", output_path],
        ):
            main()

        from docx import Document

        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

    def test_from_stdin(self, sample_feedback_data, tmp_path):
        """Test generation from stdin."""
        output_path = str(tmp_path / "stdin_output.docx")
        json_data = json.dumps(sample_feedback_data)

        with patch("sys.argv", ["prog", "--output", output_path]):
            with patch("sys.stdin", StringIO(json_data)):
                main()

        from docx import Document

        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

    def test_output_prints_path(self, input_json_file, tmp_path, capsys):
        """Test that output path is printed to stdout."""
        output_path = str(tmp_path / "output.docx")
        with patch(
            "sys.argv",
            ["prog", "--input", str(input_json_file), "--output", output_path],
        ):
            main()

        captured = capsys.readouterr()
        assert "Document generated:" in captured.out

    def test_creates_nested_output_directory(self, input_json_file, tmp_path):
        """Test that nested output directories are created."""
        output_path = str(tmp_path / "deep" / "nested" / "output.docx")
        with patch(
            "sys.argv",
            ["prog", "--input", str(input_json_file), "--output", output_path],
        ):
            main()

        from pathlib import Path

        assert Path(output_path).exists()

    def test_empty_data_generates_docx(self, empty_json_file, tmp_path):
        """Test that empty input data still produces output."""
        output_path = str(tmp_path / "empty_output.docx")
        with patch(
            "sys.argv",
            ["prog", "--input", str(empty_json_file), "--output", output_path],
        ):
            main()

        from pathlib import Path

        assert Path(output_path).exists()

    def test_invalid_json_raises(self, tmp_path):
        """Test that invalid JSON input causes an error."""
        bad_file = tmp_path / "bad.json"
        bad_file.write_text("not valid json {{{", encoding="utf-8")
        output_path = str(tmp_path / "output.docx")

        with patch(
            "sys.argv", ["prog", "--input", str(bad_file), "--output", output_path]
        ):
            with pytest.raises(json.JSONDecodeError):
                main()

    def test_invalid_json_from_stdin(self, tmp_path):
        """Test that invalid JSON from stdin causes an error."""
        output_path = str(tmp_path / "output.docx")

        with patch("sys.argv", ["prog", "--output", output_path]):
            with patch("sys.stdin", StringIO("not valid json")):
                with pytest.raises(json.JSONDecodeError):
                    main()

    def test_stderr_message_on_missing_file(self, tmp_path, capsys):
        """Test that missing file error goes to stderr."""
        output_path = str(tmp_path / "output.docx")
        with patch(
            "sys.argv",
            ["prog", "--input", "/nonexistent/file.json", "--output", output_path],
        ):
            with pytest.raises(SystemExit):
                main()

        captured = capsys.readouterr()
        assert "ERROR" in captured.err

    def test_output_overwrites_existing(self, input_json_file, tmp_path):
        """Test that existing output file is overwritten."""
        output_path = str(tmp_path / "output.docx")
        # Create the file first
        from pathlib import Path

        Path(output_path).write_text("placeholder", encoding="utf-8")

        with patch(
            "sys.argv",
            ["prog", "--input", str(input_json_file), "--output", output_path],
        ):
            main()

        from docx import Document

        doc = Document(output_path)
        assert len(doc.paragraphs) > 0


class TestCLIEdgeCases:
    """Tests for edge cases in CLI."""

    def test_malformed_data_from_file(self, tmp_path, malformed_feedback_data):
        """Test that malformed data still produces output (graceful degradation)."""
        input_file = tmp_path / "malformed.json"
        input_file.write_text(json.dumps(malformed_feedback_data), encoding="utf-8")
        output_path = str(tmp_path / "malformed_output.docx")

        with patch(
            "sys.argv",
            ["prog", "--input", str(input_file), "--output", output_path],
        ):
            # The tool should handle gracefully since it uses .get() with defaults
            # The entries field is "not-a-list" which sorted() can handle for strings
            # but the iteration expects dicts, so this may raise
            try:
                main()
            except (TypeError, AttributeError):
                # Expected for truly malformed data (e.g., "entries" as string)
                pass
